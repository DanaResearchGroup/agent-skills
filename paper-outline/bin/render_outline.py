#!/usr/bin/env python3
"""Render Outline.md into the group's Outline.docx.

The group template (assets/Outline-template.docx) is the style carrier: it is
opened, emptied of its placeholder prose, and refilled from the markdown. That
keeps the fonts, heading styles, and bullet numbering the group already uses,
so a rendered outline is indistinguishable from a hand-written one.

Markdown contract -- the shape `paper-outline` writes:

    # Title Here            -> Heading 1
    ## Subsection 1         -> Heading 2
    - a bullet              -> List Paragraph, bulleted
    anything else           -> Normal
    (blank line)            -> dropped; Word spacing comes from the styles

Usage:
    render_outline.py <Outline.md> [-o <Outline.docx>]

Defaults the output to Outline.docx beside the input.
"""

import argparse
import copy
import sys
from pathlib import Path

try:
    import docx
except ImportError:
    sys.exit("python-docx is missing: pip install python-docx")

TEMPLATE = Path(__file__).resolve().parent.parent / "assets" / "Outline-template.docx"


def _bullet_ppr(doc):
    """Deep-copy the paragraph properties of the template's first bullet.

    Carries the numPr (which numbering.xml list the bullet belongs to) and the
    indent along with it -- setting `style='List Paragraph'` alone yields an
    indented paragraph with no bullet glyph.
    """
    for p in doc.paragraphs:
        if p.style.name == "List Paragraph" and p._p.pPr is not None:
            if p._p.pPr.xpath("./w:numPr"):
                return copy.deepcopy(p._p.pPr)
    return None


def _empty_body(doc):
    """Strip every block from the body, preserving the trailing sectPr."""
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def render(md_path: Path, docx_path: Path, template: Path = TEMPLATE) -> Path:
    if not template.exists():
        sys.exit(f"template missing: {template}")
    doc = docx.Document(str(template))
    bullet_ppr = _bullet_ppr(doc)
    _empty_body(doc)

    for raw in md_path.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style="Heading 2")
        elif line.startswith("# "):
            doc.add_paragraph(line[2:].strip(), style="Heading 1")
        elif line.lstrip().startswith(("- ", "* ")):
            para = doc.add_paragraph(line.lstrip()[2:].strip())
            if bullet_ppr is None:
                para.style = "List Paragraph"
            else:
                if para._p.pPr is not None:
                    para._p.remove(para._p.pPr)
                para._p.insert(0, copy.deepcopy(bullet_ppr))
        else:
            doc.add_paragraph(line, style="Normal")

    doc.save(str(docx_path))
    return docx_path


def main():
    ap = argparse.ArgumentParser(description="Render Outline.md to the group's Outline.docx")
    ap.add_argument("markdown", type=Path, help="the Outline.md to render")
    ap.add_argument("-o", "--output", type=Path, help="output .docx (default: Outline.docx beside the input)")
    ap.add_argument("--template", type=Path, default=TEMPLATE, help="override the style-carrier template")
    args = ap.parse_args()

    if not args.markdown.exists():
        sys.exit(f"no such file: {args.markdown}")
    out = args.output or args.markdown.with_name("Outline.docx")
    print(render(args.markdown, out, args.template))


if __name__ == "__main__":
    main()
