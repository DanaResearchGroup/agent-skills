#!/usr/bin/env python3
"""Round-trip tests for render_outline.py.

Guards the parts of the renderer that fail *silently* rather than loudly: a
bullet that loses its numPr still renders, just without a bullet glyph, and a
body emptied too eagerly takes the page setup with it. Both produce a valid
.docx that is quietly wrong -- exactly what a test is for.
"""

import subprocess
import sys
import tempfile
from pathlib import Path

SKILL = Path(__file__).resolve().parent.parent
RENDER = SKILL / "bin" / "render_outline.py"
TEMPLATE = SKILL / "assets" / "Outline-template.docx"

try:
    import docx
except ImportError:
    sys.exit("SKIP: python-docx not installed")

RUN = FAIL = 0

SAMPLE = """\
# A Title

Scope: one sentence.

Paper Objective(s):
- first claim
- second claim

Outline:

# Introduction

Prose line.

# Results and discussion

## Subsection 1

More prose.

# Conclusions

- a conclusion
"""


def check(name, cond, detail=""):
    global RUN, FAIL
    RUN += 1
    if cond:
        print(f"\033[32mPASS\033[0m {name}")
    else:
        FAIL += 1
        print(f"\033[31mFAIL\033[0m {name}" + (f" -- {detail}" if detail else ""))


def styles(doc):
    return [(p.style.name, p.text) for p in doc.paragraphs]


def is_bullet(p):
    pr = p._p.pPr
    return pr is not None and bool(pr.xpath("./w:numPr"))


def main():
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        md = td / "Outline.md"
        md.write_text(SAMPLE, encoding="utf-8")

        proc = subprocess.run(
            [sys.executable, str(RENDER), str(md)], capture_output=True, text=True
        )
        check("renderer exits 0", proc.returncode == 0, proc.stderr.strip())

        out = td / "Outline.docx"
        check("defaults output to Outline.docx beside the input", out.exists())
        if not out.exists():
            return finish()

        doc = docx.Document(str(out))
        pairs = styles(doc)

        check("'# ' becomes Heading 1",
              any(s == "Heading 1" and t == "A Title" for s, t in pairs))
        check("'## ' becomes Heading 2",
              any(s == "Heading 2" and t == "Subsection 1" for s, t in pairs))
        check("plain text becomes Normal",
              any(s == "Normal" and t == "Prose line." for s, t in pairs))

        # The subtle one: 'List Paragraph' alone indents without a bullet glyph,
        # so the renderer deep-copies the template bullet's pPr to carry numPr.
        bullets = [p for p in doc.paragraphs if p.text in
                   ("first claim", "second claim", "a conclusion")]
        check("all three bullets survive", len(bullets) == 3, f"got {len(bullets)}")
        check("bullets carry real numbering (numPr)",
              bullets and all(is_bullet(p) for p in bullets),
              "a bullet lost its numPr -- it will render with no glyph")

        check("blank lines are dropped", all(t.strip() for _, t in pairs))

        # Emptying the body must spare the sectPr, or page setup is lost.
        check("sectPr (page setup) preserved",
              bool(doc.element.body.xpath("./w:sectPr")))

        # Template placeholder prose must not leak into a rendered outline.
        check("template placeholder prose is gone",
              not any("What tools and methods" in t for _, t in pairs))

        # Explicit -o path.
        alt = td / "Custom.docx"
        proc = subprocess.run(
            [sys.executable, str(RENDER), str(md), "-o", str(alt)],
            capture_output=True, text=True,
        )
        check("-o writes the named file", alt.exists() and proc.returncode == 0)

        # Error path: a missing input fails loudly rather than writing junk.
        proc = subprocess.run(
            [sys.executable, str(RENDER), str(td / "nope.md")],
            capture_output=True, text=True,
        )
        check("missing input exits non-zero", proc.returncode != 0)

        # The template itself is the style carrier -- guard its content.
        tdoc = docx.Document(str(TEMPLATE))
        ttext = "\n".join(p.text for p in tdoc.paragraphs)
        check("template says 'tools', not 'tolls'",
              "tolls" not in ttext and "What tools and methods" in ttext)
        for section in ("Introduction", "Methods", "Results and discussion",
                        "Conclusions"):
            check(f"template keeps the '{section}' heading",
                  any(p.style.name.startswith("Heading") and p.text.strip() == section
                      for p in tdoc.paragraphs))

    return finish()


def finish():
    print(f"\ntest-render.py: {RUN} run, {FAIL} failed")
    return 1 if FAIL else 0


if __name__ == "__main__":
    sys.exit(main())
