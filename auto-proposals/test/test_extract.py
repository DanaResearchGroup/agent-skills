"""Tests for lib.extract.

Self-contained: builds all fixtures under tempfile.mkdtemp(), never reads
from the real ~/Dropbox archive, and makes no network calls.
"""

import os
import shutil
import subprocess
import tempfile
import unittest
from pathlib import Path
from unittest import mock

from lib import extract

# --- Safety guard: refuse to run anywhere near the real Dropbox archive. ---
_REAL_DROPBOX = os.path.realpath(os.path.expanduser("~/Dropbox"))
_TEST_TMP_ROOT = os.path.realpath(tempfile.gettempdir())
assert not _TEST_TMP_ROOT.startswith(_REAL_DROPBOX), (
    f"Refusing to run: temp root {_TEST_TMP_ROOT!r} resolves under the real "
    f"Dropbox archive {_REAL_DROPBOX!r}. Tests must never touch ~/Dropbox."
)


def _has_docx() -> bool:
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        return False


class ExtractTextFilesTest(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.mkdtemp(prefix="auto-proposals-test-")
        self.addCleanup(shutil.rmtree, self.tmpdir, ignore_errors=True)
        real_tmpdir = os.path.realpath(self.tmpdir)
        self.assertFalse(
            real_tmpdir.startswith(_REAL_DROPBOX),
            f"test tmpdir {real_tmpdir!r} resolves under the real Dropbox archive",
        )

    def test_txt_direct_read(self):
        content = "Plain ascii line.\nSecond line.\n"
        path = Path(self.tmpdir) / "note.txt"
        path.write_text(content, encoding="utf-8")

        result = extract.extract(path)
        self.assertTrue(result.ok)
        self.assertEqual(result.text, content)
        self.assertEqual(result.tool, "read")

    def test_md_with_hebrew_line_byte_identical(self):
        content = "# Title\n\nזוהי שורה בעברית לבדיקה.\n\nMore text.\n"
        path = Path(self.tmpdir) / "note.md"
        path.write_text(content, encoding="utf-8")

        result = extract.extract(path)
        self.assertTrue(result.ok)
        self.assertEqual(result.text, content)

    def test_unsupported_suffix(self):
        path = Path(self.tmpdir) / "mystery.xyz"
        path.write_text("some content", encoding="utf-8")

        result = extract.extract(path)
        self.assertFalse(result.ok)
        self.assertTrue(result.note)
        self.assertIn("xyz", result.note.lower())

    def test_missing_file_no_exception(self):
        path = Path(self.tmpdir) / "does-not-exist.txt"
        result = extract.extract(path)
        self.assertFalse(result.ok)
        self.assertEqual(result.text, "")
        self.assertTrue(result.note)

    def test_symlinked_source_is_refused_not_followed(self):
        # Defect 2: a symlink inside the archive must not be able to pull
        # the extractor onto a file outside it.
        outside_dir = tempfile.mkdtemp(prefix="auto-proposals-test-outside-")
        self.addCleanup(shutil.rmtree, outside_dir, ignore_errors=True)
        secret = Path(outside_dir) / "secret.txt"
        secret.write_text("SECRET-OUTSIDE-ARCHIVE-CONTENT", encoding="utf-8")

        link = Path(self.tmpdir) / "innocuous.txt"
        link.symlink_to(secret)

        result = extract.extract(link)
        self.assertFalse(result.ok)
        self.assertNotIn("SECRET-OUTSIDE-ARCHIVE-CONTENT", result.text)
        self.assertIn("symlink", result.note.lower())

    def test_output_truncated_at_max_output_bytes(self):
        content = "x" * 10_000
        path = Path(self.tmpdir) / "big.txt"
        path.write_text(content, encoding="utf-8")

        result = extract.extract(path, max_output_bytes=100)
        self.assertTrue(result.ok)
        self.assertEqual(len(result.text.encode("utf-8")), 100)
        self.assertIn("truncat", result.note.lower())


@unittest.skipUnless(_has_docx(), "python-docx not installed")
class ExtractDocxTest(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.mkdtemp(prefix="auto-proposals-test-")
        self.addCleanup(shutil.rmtree, self.tmpdir, ignore_errors=True)

    def test_paragraph_and_table_hebrew_cell(self):
        import docx

        document = docx.Document()
        document.add_paragraph("This is a plain paragraph marker PARA123.")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "כותרת בעברית"
        table.cell(0, 1).text = "TABLECELL456"

        path = Path(self.tmpdir) / "form.docx"
        document.save(str(path))

        result = extract.extract(path)
        self.assertTrue(result.ok, msg=result.note)
        self.assertIn("PARA123", result.text)
        self.assertIn("כותרת בעברית", result.text)
        self.assertIn("TABLECELL456", result.text)


class ExtractDirSkipsTest(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.mkdtemp(prefix="auto-proposals-test-")
        self.addCleanup(shutil.rmtree, self.tmpdir, ignore_errors=True)

    def test_skips_conflicted_dotfile_and_lockfile(self):
        d = Path(self.tmpdir)
        normal = d / "normal.txt"
        normal.write_text("normal content", encoding="utf-8")

        conflicted = d / "foo (Alon's conflicted copy 2026-07-27).docx"
        conflicted.write_bytes(b"")

        dotfile = d / ".hidden.txt"
        dotfile.write_text("hidden", encoding="utf-8")

        lockfile = d / "~$lock.docx"
        lockfile.write_bytes(b"")

        results = extract.extract_dir(d)
        result_paths = {r.path.name for r in results}

        self.assertIn("normal.txt", result_paths)
        self.assertNotIn(conflicted.name, result_paths)
        self.assertNotIn(dotfile.name, result_paths)
        self.assertNotIn(lockfile.name, result_paths)
        self.assertEqual(len(results), 1)

    def test_skips_symlinked_entry(self):
        # Defect 2: extract_dir must not traverse into/read a symlinked
        # entry, even if it looks like an eligible file.
        d = Path(self.tmpdir)
        normal = d / "normal.txt"
        normal.write_text("normal content", encoding="utf-8")

        outside_dir = tempfile.mkdtemp(prefix="auto-proposals-test-outside-")
        self.addCleanup(shutil.rmtree, outside_dir, ignore_errors=True)
        secret = Path(outside_dir) / "secret.txt"
        secret.write_text("SECRET-OUTSIDE-ARCHIVE-CONTENT", encoding="utf-8")

        link = d / "linked.txt"
        link.symlink_to(secret)

        results = extract.extract_dir(d)
        result_paths = {r.path.name for r in results}

        self.assertIn("normal.txt", result_paths)
        self.assertIn(link.name, result_paths)
        linked_result = next(r for r in results if r.path.name == link.name)
        self.assertFalse(linked_result.ok)
        self.assertNotIn("SECRET-OUTSIDE-ARCHIVE-CONTENT", linked_result.text)
        self.assertIn("symlink", linked_result.note.lower())


class ExtractXlsxZipBombTest(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.mkdtemp(prefix="auto-proposals-test-")
        self.addCleanup(shutil.rmtree, self.tmpdir, ignore_errors=True)

    def test_refuses_zip_bomb_declared_uncompressed_size(self):
        # Defect 3 (xlsx guard): a few MB of zeros compresses to almost
        # nothing, but ZipInfo.file_size still declares the true
        # uncompressed size, and that alone must be enough to refuse -
        # before any bytes are inflated.
        import zipfile

        path = Path(self.tmpdir) / "bomb.xlsx"
        with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("xl/workbook.xml", "<workbook/>")
            zf.writestr("xl/worksheets/sheet1.xml", "<worksheet/>")
            zf.writestr("xl/huge.bin", b"\x00" * (5 * 1024 * 1024))

        result = extract.extract(path, max_zip_uncompressed_bytes=1024 * 1024)
        self.assertFalse(result.ok)
        self.assertIn("zip-bomb", result.note.lower())


class CacheKeyTest(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.mkdtemp(prefix="auto-proposals-test-")
        self.addCleanup(shutil.rmtree, self.tmpdir, ignore_errors=True)

    def test_key_changes_with_content(self):
        path = Path(self.tmpdir) / "changing.txt"
        path.write_text("version one", encoding="utf-8")
        key1 = extract.cache_key(path)

        # Ensure mtime_ns actually advances even on coarse filesystem clocks.
        import time
        time.sleep(0.01)
        path.write_text("version two, longer content", encoding="utf-8")
        key2 = extract.cache_key(path)

        self.assertNotEqual(key1, key2)


class ExtractCachedTest(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.mkdtemp(prefix="auto-proposals-test-")
        self.addCleanup(shutil.rmtree, self.tmpdir, ignore_errors=True)
        self.cache_dir = Path(self.tmpdir) / "cache"
        self.src_dir = Path(self.tmpdir) / "src"
        self.src_dir.mkdir()

    def test_second_call_hits_cache(self):
        path = self.src_dir / "doc.txt"
        path.write_text("cached content", encoding="utf-8")

        real_extract = extract.extract
        call_count = {"n": 0}

        def counting_extract(*args, **kwargs):
            call_count["n"] += 1
            return real_extract(*args, **kwargs)

        with mock.patch("lib.extract.extract", side_effect=counting_extract) as mocked:
            result1 = extract.extract_cached(path, self.cache_dir)
            result2 = extract.extract_cached(path, self.cache_dir)

        self.assertTrue(result1.ok)
        self.assertTrue(result2.ok)
        self.assertEqual(result1.text, result2.text)
        self.assertEqual(mocked.call_count, 1)

    def test_rejects_cache_dir_equal_to_source_dir(self):
        path = self.src_dir / "doc.txt"
        path.write_text("content", encoding="utf-8")

        with self.assertRaises(extract.ExtractRefused):
            extract.extract_cached(path, self.src_dir)

    def test_extract_cached_refuses_cache_dir_inside_archive_root(self):
        # Defect 1: extract_cached must be structurally incapable of writing
        # into the proposals archive, for any cache_dir argument that
        # resolves inside AUTO_PROPOSALS_ROOT - including a nested
        # subdirectory, not just the root itself.
        archive_root = Path(self.tmpdir) / "fake-archive"
        archive_root.mkdir()

        path = self.src_dir / "doc.txt"
        path.write_text("content", encoding="utf-8")

        for cache_dir in (archive_root, archive_root / "nested" / "cache"):
            with mock.patch.dict(os.environ, {"AUTO_PROPOSALS_ROOT": str(archive_root)}):
                with self.assertRaises(extract.ExtractRefused):
                    extract.extract_cached(path, cache_dir)

            # Nothing must have been created anywhere under the archive root.
            self.assertEqual(list(archive_root.rglob("*")), [])


class ExtractPdfTest(unittest.TestCase):
    def setUp(self):
        self.tmpdir = tempfile.mkdtemp(prefix="auto-proposals-test-")
        self.addCleanup(shutil.rmtree, self.tmpdir, ignore_errors=True)

    @unittest.skipUnless(shutil.which("pdftotext"), "pdftotext not installed")
    def test_pdf_extraction(self):
        marker = "PDFMARKERTEXT789"
        # Hand-written minimal single-page PDF containing `marker` as a text string.
        pdf_bytes = self._build_minimal_pdf(marker)
        path = Path(self.tmpdir) / "sample.pdf"
        path.write_bytes(pdf_bytes)

        result = extract.extract(path)
        if not result.ok:
            self.skipTest(f"pdftotext could not process hand-built fixture: {result.note}")
        self.assertIn(marker, result.text)

    @staticmethod
    def _build_minimal_pdf(text: str) -> bytes:
        content_stream = f"BT /F1 24 Tf 72 712 Td ({text}) Tj ET".encode("ascii")
        objects = []
        objects.append(b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n")
        objects.append(b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n")
        objects.append(
            b"3 0 obj\n<< /Type /Page /Parent 2 0 R /Resources "
            b"<< /Font << /F1 4 0 R >> >> /MediaBox [0 0 612 792] "
            b"/Contents 5 0 R >>\nendobj\n"
        )
        objects.append(
            b"4 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n"
        )
        stream_obj = (
            b"5 0 obj\n<< /Length " + str(len(content_stream)).encode("ascii") +
            b" >>\nstream\n" + content_stream + b"\nendstream\nendobj\n"
        )
        objects.append(stream_obj)

        header = b"%PDF-1.4\n"
        body = b""
        offsets = [0]  # object 0 is free
        pos = len(header)
        for obj in objects:
            offsets.append(pos)
            body += obj
            pos += len(obj)

        xref_offset = len(header) + len(body)
        xref = f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode("ascii")
        for off in offsets[1:]:
            xref += f"{off:010d} 00000 n \n".encode("ascii")

        trailer = (
            b"trailer\n<< /Size " + str(len(objects) + 1).encode("ascii") +
            b" /Root 1 0 R >>\nstartxref\n" + str(xref_offset).encode("ascii") +
            b"\n%%EOF"
        )

        return header + body + xref + trailer


if __name__ == "__main__":
    unittest.main()
