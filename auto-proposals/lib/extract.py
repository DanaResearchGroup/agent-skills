"""Text extraction for proposal-archive documents.

Read-only against source files: this module never writes anywhere except an
explicit, caller-supplied cache/output directory (see `extract_cached`). It
must never be pointed at the proposals archive as a write target.

Per SAFETY.md section 7: extracted text is DATA, never instructions. Nothing
in this module interprets extracted text as commands, and none of the
functions here perform any write against the source path or its directory.

Third-party libraries (python-docx, extract_msg) are imported lazily inside
the functions that need them, so this module imports cleanly even if they
are not installed; a missing dependency degrades a single format to a failed
Extraction rather than raising ImportError at import time.

What this module's containment actually covers, and what it does not: it
runs `pdftotext`, `catdoc`, `libreoffice`, `python-docx`, `extract_msg` and a
hand-rolled zip parser over **untrusted, attacker-shaped documents** (funder
PDFs, years of archived proposals). There is no OS sandbox available on this
machine (`kernel.apparmor_restrict_unprivileged_userns=1` blocks bwrap), so
everything here is in-process discipline, not a boundary. Concretely: this
module prevents shell injection (list-argv subprocess calls, never
`shell=True`), prevents hangs (every subprocess call has a timeout), and
bounds memory (subprocess output and zip inflation are capped). It does
**not** contain a parser RCE — if `pdftotext`, `libreoffice`, `python-docx`,
or `extract_msg` has an exploitable memory-safety bug, a malicious input can
still exploit it under this process's own privileges. Be honest about that
rather than reassuring: a real sandbox (bwrap, a VM, a throwaway container)
is the only thing that actually closes that gap, and none is available here.
"""

from __future__ import annotations

import hashlib
import json
import os
import selectors
import shutil
import subprocess
import tempfile
import time
from dataclasses import dataclass, asdict
from pathlib import Path

# Roughly the point below which a "successful" primary extraction is
# considered suspicious enough to try the libreoffice fallback instead.
_MIN_CHARS_BEFORE_FALLBACK = 200

_CONFLICTED_COPY_MARKERS = (
    "conflicted copy",
    "conflicted-copy",
)

# Default cap on recovered text per extraction, and on subprocess stdout
# while it is being read. See the module docstring: this bounds memory, it
# does not sandbox the parser.
_DEFAULT_MAX_OUTPUT_BYTES = 8 * 1024 * 1024  # 8 MB

# Zip-bomb guard for .xlsx parsing: reject an archive whose declared
# (pre-decompression) uncompressed size is absurd, or whose per-entry
# compression ratio is absurd, before any bytes are inflated.
_DEFAULT_MAX_ZIP_UNCOMPRESSED_BYTES = 200 * 1024 * 1024  # 200 MB
_MAX_ZIP_COMPRESSION_RATIO = 100  # declared-uncompressed / on-disk

# stderr is only ever used for a truncated error message, so it gets a
# small fixed cap regardless of what the caller asked for on stdout.
_STDERR_CAP_BYTES = 64 * 1024


class ExtractRefused(Exception):
    """Raised whenever this module is asked to do something that risks
    writing into the proposals archive (see `extract_cached`). This is a
    real exception, not an `assert` — asserts are stripped under `python
    -O`, and a safety check that can be silently compiled away is not a
    safety check.
    """


@dataclass
class Extraction:
    path: Path          # source file
    text: str            # recovered plain text ("" on failure)
    tool: str            # which extractor actually produced it
    ok: bool
    note: str            # why it failed, or a caveat (e.g. "used libreoffice fallback")
    chars: int


@dataclass
class _RunResult:
    stdout: str
    stderr: str
    returncode: int
    truncated: bool      # True if stdout was cut off at max_output_bytes


def _make_failed(path: Path, tool: str, note: str) -> Extraction:
    return Extraction(path=path, text="", tool=tool, ok=False, note=note, chars=0)


def _make_ok(path: Path, tool: str, text: str, note: str = "") -> Extraction:
    return Extraction(path=path, text=text, tool=tool, ok=True, note=note, chars=len(text))


def _cap_text(text: str, max_output_bytes: int) -> tuple[str, bool]:
    """Cap `text` at `max_output_bytes` UTF-8 bytes, truncating on a
    character boundary. Returns (possibly-truncated text, was_truncated)."""
    encoded = text.encode("utf-8")
    if len(encoded) <= max_output_bytes:
        return text, False
    return encoded[:max_output_bytes].decode("utf-8", errors="ignore"), True


def _run(
    argv: list[str],
    *,
    timeout: int,
    max_output_bytes: int | None = None,
) -> _RunResult | None:
    """Run a subprocess safely; return None (never raises) on failure to
    launch or on timeout.

    Reads stdout/stderr incrementally via a selector rather than
    `subprocess.run(capture_output=True)`: the latter buffers the child's
    entire output in memory before any cap can be applied, so a zip bomb or
    a pathological PDF could exhaust memory well before hitting the
    wall-clock timeout. Here stdout is capped at `max_output_bytes` (if
    given) — once the cap is exceeded, the child is killed immediately
    rather than left to keep producing output we would discard anyway.
    stderr is always capped at a small fixed size, since it is only ever
    used for a truncated error message.
    """
    try:
        proc = subprocess.Popen(argv, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    except FileNotFoundError:
        return None

    cap = max_output_bytes if max_output_bytes is not None else -1
    stdout_chunks: list[bytes] = []
    stderr_chunks: list[bytes] = []
    stdout_total = 0
    stderr_total = 0
    truncated = False
    timed_out = False

    sel = selectors.DefaultSelector()
    sel.register(proc.stdout, selectors.EVENT_READ, "stdout")
    sel.register(proc.stderr, selectors.EVENT_READ, "stderr")
    open_streams = 2
    deadline = time.monotonic() + timeout

    try:
        while open_streams > 0:
            remaining = deadline - time.monotonic()
            if remaining <= 0:
                timed_out = True
                break
            for key, _ in sel.select(timeout=remaining):
                chunk = os.read(key.fileobj.fileno(), 65536)
                if not chunk:
                    sel.unregister(key.fileobj)
                    open_streams -= 1
                    continue
                if key.data == "stdout":
                    if cap >= 0 and stdout_total >= cap:
                        truncated = True
                        continue
                    stdout_chunks.append(chunk)
                    stdout_total += len(chunk)
                    if cap >= 0 and stdout_total > cap:
                        truncated = True
                else:
                    if stderr_total < _STDERR_CAP_BYTES:
                        stderr_chunks.append(chunk)
                        stderr_total += len(chunk)
            if truncated:
                # Nothing more to gain from this process; stop reading and
                # kill it below rather than let it keep running/producing.
                break
    finally:
        sel.close()

    if truncated or timed_out:
        try:
            proc.kill()
        except OSError:
            pass

    try:
        proc.wait(timeout=5)
    except subprocess.TimeoutExpired:
        timed_out = True
    except OSError:
        pass

    try:
        proc.stdout.close()
    except OSError:
        pass
    try:
        proc.stderr.close()
    except OSError:
        pass

    if timed_out:
        return None

    stdout_bytes = b"".join(stdout_chunks)
    if cap >= 0 and len(stdout_bytes) > cap:
        stdout_bytes = stdout_bytes[:cap]
    stderr_bytes = b"".join(stderr_chunks)

    return _RunResult(
        stdout=stdout_bytes.decode("utf-8", errors="replace"),
        stderr=stderr_bytes.decode("utf-8", errors="replace"),
        returncode=proc.returncode if proc.returncode is not None else -1,
        truncated=truncated,
    )


def _lo_tmp_base() -> str:
    """Base directory for libreoffice fallback tmpdirs.

    Deliberately NOT the bare system /tmp: on this machine libreoffice ships
    as a snap, and snap confinement gives the snap process its own private
    view of /tmp — files it writes there under an arbitrary --outdir are
    invisible to the calling (unconfined) process, even though libreoffice
    reports success. A directory under the user's home is reliably visible
    to both sides, so anchor mkdtemp() there instead.
    """
    base = os.path.join(os.path.expanduser("~"), ".cache", "auto-proposals", "lo-tmp")
    os.makedirs(base, exist_ok=True)
    return base


def _libreoffice_fallback(path: Path, *, timeout: int, max_output_bytes: int) -> Extraction:
    """Convert `path` to plain text via headless libreoffice, in an isolated tmpdir.

    The tmpdir (which holds both the private `UserInstallation` profile and
    the converted output file) is always removed before returning, on every
    path — success, tool failure, missing output, or timeout — so a failed
    or killed fallback never leaks a stray directory under ~/.cache.
    """
    tmpdir = tempfile.mkdtemp(prefix="auto-proposals-lo-", dir=_lo_tmp_base())
    try:
        profile_dir = Path(tmpdir) / "lo"
        argv = [
            "libreoffice",
            "--headless",
            f"-env:UserInstallation=file://{profile_dir}",
            "--convert-to",
            "txt:Text",
            "--outdir",
            tmpdir,
            str(path),
        ]
        proc = _run(argv, timeout=timeout, max_output_bytes=max_output_bytes)
        if proc is None:
            return _make_failed(path, "libreoffice", "libreoffice not installed or timed out")
        if proc.returncode != 0:
            return _make_failed(
                path, "libreoffice",
                f"libreoffice fallback failed (rc={proc.returncode}): {proc.stderr.strip()[:500]}",
            )
        out_path = Path(tmpdir) / (path.stem + ".txt")
        if not out_path.exists():
            return _make_failed(path, "libreoffice", "libreoffice fallback produced no output file")
        try:
            text = out_path.read_text(encoding="utf-8", errors="replace")
        except OSError as exc:
            return _make_failed(path, "libreoffice", f"could not read libreoffice output: {exc}")
        text, was_truncated = _cap_text(text, max_output_bytes)
        note = "used libreoffice fallback"
        if was_truncated:
            note += f"; output truncated at max_output_bytes={max_output_bytes}"
        return _make_ok(path, "libreoffice", text, note=note)
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


def _maybe_fallback(
    primary: Extraction, path: Path, *, timeout: int, max_output_bytes: int,
) -> Extraction:
    """Given a primary-tool Extraction, fall back to libreoffice if it failed or looks thin."""
    if primary.ok and primary.chars >= _MIN_CHARS_BEFORE_FALLBACK:
        return primary
    fallback = _libreoffice_fallback(path, timeout=timeout, max_output_bytes=max_output_bytes)
    if fallback.ok:
        reason = "primary tool failed" if not primary.ok else "primary tool yielded too little text"
        fallback.note = f"{reason} ({primary.note or primary.tool}); used libreoffice fallback"
        return fallback
    # Both failed; prefer returning the primary's result (it may have partial
    # text) but note that fallback was also attempted and failed. A failing
    # fallback must never discard a usable primary result.
    if primary.ok:
        primary.note = (primary.note + "; " if primary.note else "") + \
            f"libreoffice fallback also attempted and failed: {fallback.note}"
        return primary
    return _make_failed(
        path, primary.tool,
        f"primary failed ({primary.note}); libreoffice fallback also failed ({fallback.note})",
    )


def _extract_pdf(path: Path, *, timeout: int, max_output_bytes: int) -> Extraction:
    proc = _run(
        ["pdftotext", "-layout", "-enc", "UTF-8", str(path), "-"],
        timeout=timeout, max_output_bytes=max_output_bytes,
    )
    if proc is None:
        primary = _make_failed(path, "pdftotext", "pdftotext not installed or timed out")
    elif proc.truncated:
        primary = _make_ok(
            path, "pdftotext", proc.stdout,
            note=f"output truncated at max_output_bytes={max_output_bytes}",
        )
    elif proc.returncode != 0:
        primary = _make_failed(path, "pdftotext", f"pdftotext failed (rc={proc.returncode}): {proc.stderr.strip()[:500]}")
    else:
        primary = _make_ok(path, "pdftotext", proc.stdout)
    return _maybe_fallback(primary, path, timeout=timeout, max_output_bytes=max_output_bytes)


def _extract_doc(path: Path, *, timeout: int, max_output_bytes: int) -> Extraction:
    proc = _run(["catdoc", str(path)], timeout=timeout, max_output_bytes=max_output_bytes)
    if proc is None:
        primary = _make_failed(path, "catdoc", "catdoc not installed or timed out")
    elif proc.truncated:
        primary = _make_ok(
            path, "catdoc", proc.stdout,
            note=f"output truncated at max_output_bytes={max_output_bytes}",
        )
    elif proc.returncode != 0:
        primary = _make_failed(path, "catdoc", f"catdoc failed (rc={proc.returncode}): {proc.stderr.strip()[:500]}")
    else:
        primary = _make_ok(path, "catdoc", proc.stdout)
    return _maybe_fallback(primary, path, timeout=timeout, max_output_bytes=max_output_bytes)


def _iter_table_cells(table):
    """Yield cell text for a table, recursing into nested tables, in row order."""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if para.text:
                    yield para.text
            for nested in cell.tables:
                yield from _iter_table_cells(nested)


def _docx_worker(path_str: str, max_output_bytes: int, result_queue) -> None:
    """Runs in a child process (see `_run_in_process_with_timeout`). Must not
    touch any shared state - only argv-in, result-via-queue-out.

    Any parse failure is caught and put on the queue as a ("fail", ...)
    tuple rather than left to crash the child silently: an uncaught
    exception here would leave the queue empty, which the parent cannot
    distinguish from a timeout, misreporting a real parse error as a hang.
    """
    try:
        import docx  # python-docx; lazy import, only ever runs inside the child

        path = Path(path_str)
        document = docx.Document(str(path))
        paragraphs = [p.text for p in document.paragraphs if p.text]
        table_lines: list[str] = []
        for table in document.tables:
            table_lines.extend(_iter_table_cells(table))
        parts = list(paragraphs)
        if table_lines:
            parts.append("")
            parts.append(
                "[Table content below; order is 'all paragraphs, then all tables' — "
                "not necessarily original document order]"
            )
            parts.extend(table_lines)
        text = "\n".join(parts)
        text, was_truncated = _cap_text(text, max_output_bytes)
        note = (
            "" if not table_lines else
            "paragraphs and tables both extracted; tables emitted after all "
            "paragraphs (not interleaved in original document order)"
        )
        if was_truncated:
            trunc_note = f"output truncated at max_output_bytes={max_output_bytes}"
            note = f"{note}; {trunc_note}" if note else trunc_note
        result_queue.put(("ok", text, note))
    except Exception as exc:  # noqa: BLE001 - report, never let the child crash silently
        result_queue.put(("fail", str(exc)))


def _run_in_process_with_timeout(worker, args: tuple, *, timeout: int):
    """Run `worker(*args, result_queue)` in a `multiprocessing` child and
    wait up to `timeout` seconds for it to put a result on the queue.

    `python-docx` and `extract_msg` are called in-process (unlike the
    subprocess-based extractors above), so they have no wall-clock timeout of
    their own: a pathological file can hang the whole run, which under cron
    means a silently stalled weekly job. Running the actual parse in a child
    process gives us the same hard timeout the subprocess extractors get -
    on expiry the child is terminated (and killed if it doesn't die
    promptly) and this returns `None`, same contract as `_run()` above.

    Returns `None` on timeout, on the child dying without a result (crash),
    or on any other failure to get a clean result back; otherwise returns
    the tuple the worker put on the queue.
    """
    import multiprocessing

    # "fork" (the Linux default) rather than "spawn": besides being cheaper
    # (no re-import of this process's modules), it means a test that
    # monkeypatches e.g. `docx.Document` in the parent before calling this
    # sees that patch honoured in the child too, since fork copies the
    # parent's already-imported module state instead of re-importing fresh.
    ctx = multiprocessing.get_context("fork")
    result_queue = ctx.Queue()
    proc = ctx.Process(target=worker, args=(*args, result_queue))
    proc.start()
    proc.join(timeout)

    if proc.is_alive():
        proc.terminate()
        proc.join(5)
        if proc.is_alive():
            proc.kill()
            proc.join()
        result_queue.close()
        return None

    try:
        result = result_queue.get_nowait()
    except Exception:  # noqa: BLE001 - empty queue, corrupt pickle, etc.
        result = None
    result_queue.close()
    return result


def _extract_docx(path: Path, *, timeout: int, max_output_bytes: int) -> Extraction:
    try:
        import docx  # noqa: F401 - probe only, to fail fast with a clear note
    except ImportError:
        primary = _make_failed(path, "python-docx", "python-docx not installed")
        return _maybe_fallback(primary, path, timeout=timeout, max_output_bytes=max_output_bytes)

    try:
        result = _run_in_process_with_timeout(
            _docx_worker, (str(path), max_output_bytes), timeout=timeout,
        )
        if result is None:
            primary = _make_failed(
                path, "python-docx",
                f"python-docx timed out after {timeout}s (in-process extraction watchdog)",
            )
        elif result[0] == "ok":
            _, text, note = result
            primary = _make_ok(path, "python-docx", text, note=note)
        else:
            _, err = result
            primary = _make_failed(path, "python-docx", f"python-docx failed to parse: {err}")
    except Exception as exc:  # noqa: BLE001 - convert any docx parse failure into a failed Extraction
        primary = _make_failed(path, "python-docx", f"python-docx failed to parse: {exc}")

    return _maybe_fallback(primary, path, timeout=timeout, max_output_bytes=max_output_bytes)


def _msg_worker(path_str: str, max_output_bytes: int, result_queue) -> None:
    """Runs in a child process (see `_run_in_process_with_timeout`).

    Any parse failure is caught and put on the queue as a ("fail", ...)
    tuple rather than left to crash the child silently: an uncaught
    exception here would leave the queue empty, which the parent cannot
    distinguish from a timeout, misreporting a real parse error as a hang.
    """
    try:
        import extract_msg  # lazy import, only ever runs inside the child

        path = Path(path_str)
        msg = extract_msg.Message(str(path))
        try:
            subject = msg.subject or ""
            sender = msg.sender or ""
            to = msg.to or ""
            date = str(msg.date) if msg.date else ""
            body = msg.body or ""
            attachments = []
            for att in getattr(msg, "attachments", []) or []:
                name = getattr(att, "longFilename", None) or getattr(att, "shortFilename", None)
                if name:
                    attachments.append(str(name))
            parts = [
                f"Subject: {subject}",
                f"From: {sender}",
                f"To: {to}",
                f"Date: {date}",
                "",
                body,
            ]
            if attachments:
                parts.append("")
                parts.append("Attachments:")
                parts.extend(f"- {name}" for name in attachments)
            text = "\n".join(parts)
            text, was_truncated = _cap_text(text, max_output_bytes)
            note = f"output truncated at max_output_bytes={max_output_bytes}" if was_truncated else ""
            result_queue.put(("ok", text, note))
        finally:
            try:
                msg.close()
            except Exception:  # noqa: BLE001
                pass
    except Exception as exc:  # noqa: BLE001 - report, never let the child crash silently
        result_queue.put(("fail", str(exc)))


def _extract_msg(path: Path, *, timeout: int, max_output_bytes: int) -> Extraction:
    try:
        import extract_msg  # noqa: F401 - probe only, to fail fast with a clear note
    except ImportError:
        return _make_failed(path, "extract_msg", "extract_msg not installed")

    try:
        result = _run_in_process_with_timeout(
            _msg_worker, (str(path), max_output_bytes), timeout=timeout,
        )
    except Exception as exc:  # noqa: BLE001
        return _make_failed(path, "extract_msg", f"extract_msg failed to parse: {exc}")

    if result is None:
        return _make_failed(
            path, "extract_msg",
            f"extract_msg timed out after {timeout}s (in-process extraction watchdog)",
        )
    if result[0] == "ok":
        _, text, note = result
        return _make_ok(path, "extract_msg", text, note=note)
    _, err = result
    return _make_failed(path, "extract_msg", f"extract_msg failed to parse: {err}")


def _extract_xlsx(
    path: Path, *, timeout: int, max_output_bytes: int, max_zip_uncompressed_bytes: int,
) -> Extraction:
    """Best-effort .xlsx text extraction via zipfile + shared strings XML.

    No new dependency: reads xl/sharedStrings.xml and xl/worksheets/sheet*.xml
    directly out of the zip. This is deliberately best-effort — it recovers
    shared-string cell text but not formulas, numeric-only cells, or richly
    formatted content. Considered acceptable for a first pass; flagged as
    fragile in reporting.

    Before touching any entry's compressed bytes, this refuses the archive
    outright if `ZipInfo.file_size` (the declared, pre-decompression
    uncompressed size) totals more than `max_zip_uncompressed_bytes`, or if
    any entry's declared compression ratio is absurd — the classic zip-bomb
    guard, checked from metadata alone so it costs nothing to apply even to
    a hostile file.
    """
    import zipfile
    import re

    try:
        with zipfile.ZipFile(path) as zf:
            infos = zf.infolist()

            total_uncompressed = sum(info.file_size for info in infos)
            if total_uncompressed > max_zip_uncompressed_bytes:
                return _make_failed(
                    path, "xlsx-zipfile",
                    f"refused: declared uncompressed size {total_uncompressed} bytes exceeds "
                    f"cap {max_zip_uncompressed_bytes} bytes (zip-bomb guard)",
                )
            for info in infos:
                if info.compress_size > 0:
                    ratio = info.file_size / info.compress_size
                    if ratio > _MAX_ZIP_COMPRESSION_RATIO and info.file_size > 1024 * 1024:
                        return _make_failed(
                            path, "xlsx-zipfile",
                            f"refused: entry {info.filename!r} has compression ratio "
                            f"{ratio:.0f}x (declared {info.file_size} bytes from "
                            f"{info.compress_size} bytes on disk), exceeding cap "
                            f"{_MAX_ZIP_COMPRESSION_RATIO}x (zip-bomb guard)",
                        )

            names = [info.filename for info in infos]
            shared_strings: list[str] = []
            if "xl/sharedStrings.xml" in names:
                raw = zf.read("xl/sharedStrings.xml").decode("utf-8", errors="replace")
                # Each <si> entry may contain one or more <t> text runs.
                for si in re.findall(r"<si\b.*?</si>", raw, flags=re.S):
                    texts = re.findall(r"<t[^>]*>(.*?)</t>", si, flags=re.S)
                    joined = "".join(texts)
                    joined = (
                        joined.replace("&amp;", "&")
                        .replace("&lt;", "<")
                        .replace("&gt;", ">")
                        .replace("&quot;", '"')
                        .replace("&apos;", "'")
                    )
                    shared_strings.append(joined)

            sheet_names = sorted(n for n in names if n.startswith("xl/worksheets/sheet") and n.endswith(".xml"))
            lines: list[str] = []
            for sheet_name in sheet_names:
                raw = zf.read(sheet_name).decode("utf-8", errors="replace")
                lines.append(f"[{sheet_name}]")
                for row_match in re.findall(r"<row\b.*?</row>", raw, flags=re.S):
                    cell_texts = []
                    for cell_match in re.findall(r"<c\b[^>]*>.*?</c>|<c\b[^>]*/>", row_match, flags=re.S):
                        is_shared = 't="s"' in cell_match
                        v_match = re.search(r"<v>(.*?)</v>", cell_match, flags=re.S)
                        if v_match is None:
                            continue
                        val = v_match.group(1)
                        if is_shared:
                            try:
                                idx = int(val)
                                val = shared_strings[idx] if 0 <= idx < len(shared_strings) else val
                            except ValueError:
                                pass
                        cell_texts.append(val)
                    if cell_texts:
                        lines.append("\t".join(cell_texts))

            text = "\n".join(lines)
            if not text.strip():
                return _make_failed(path, "xlsx-zipfile", "no extractable text found (empty or unsupported layout)")
            text, was_truncated = _cap_text(text, max_output_bytes)
            note = "best-effort extraction via zipfile+sharedStrings; formulas/numeric-only cells may be missing"
            if was_truncated:
                note += f"; output truncated at max_output_bytes={max_output_bytes}"
            return _make_ok(path, "xlsx-zipfile", text, note=note)
    except (zipfile.BadZipFile, KeyError, OSError) as exc:
        return _make_failed(path, "xlsx-zipfile", f"xlsx best-effort extraction failed: {exc}")


def _extract_text_file(path: Path, *, timeout: int, max_output_bytes: int) -> Extraction:
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
    except OSError as exc:
        return _make_failed(path, "read", f"could not read file: {exc}")
    text, was_truncated = _cap_text(text, max_output_bytes)
    note = f"output truncated at max_output_bytes={max_output_bytes}" if was_truncated else ""
    return _make_ok(path, "read", text, note=note)


def extract(
    path: Path,
    *,
    timeout: int = 120,
    max_output_bytes: int = _DEFAULT_MAX_OUTPUT_BYTES,
    max_zip_uncompressed_bytes: int = _DEFAULT_MAX_ZIP_UNCOMPRESSED_BYTES,
) -> Extraction:
    """Extract plain text from a single file, dispatching on suffix.

    Never raises: any tool failure, missing dependency, timeout, or missing
    file is converted into a failed Extraction (ok=False) with a note.

    Refuses a symlinked source outright rather than following it: `path` is
    checked with `Path.is_symlink()` (an `os.lstat`-based check that never
    follows the link) before anything else, so a symlink planted inside the
    archive cannot be used to pull this function onto a file outside it.

    Recovered text is capped at `max_output_bytes`; a truncated extraction
    is still `ok=True` (it is real, just incomplete), with a `note` saying
    so, so callers never mistake it for a complete extraction.
    """
    path = Path(path)
    if path.is_symlink():
        return _make_failed(path, "none", f"refusing to follow symlinked source: {path}")
    if not path.exists():
        return _make_failed(path, "none", f"file does not exist: {path}")

    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            return _extract_pdf(path, timeout=timeout, max_output_bytes=max_output_bytes)
        if suffix == ".docx":
            return _extract_docx(path, timeout=timeout, max_output_bytes=max_output_bytes)
        if suffix == ".doc":
            return _extract_doc(path, timeout=timeout, max_output_bytes=max_output_bytes)
        if suffix == ".msg":
            return _extract_msg(path, timeout=timeout, max_output_bytes=max_output_bytes)
        if suffix == ".xlsx":
            return _extract_xlsx(
                path, timeout=timeout, max_output_bytes=max_output_bytes,
                max_zip_uncompressed_bytes=max_zip_uncompressed_bytes,
            )
        if suffix in (".txt", ".md"):
            return _extract_text_file(path, timeout=timeout, max_output_bytes=max_output_bytes)
        return _make_failed(path, "none", f"unsupported file suffix: {suffix!r}")
    except Exception as exc:  # noqa: BLE001 - last-resort guard, extract() must never raise
        return _make_failed(path, "unknown", f"unexpected error during extraction: {exc}")


def _is_conflicted_copy(name: str) -> bool:
    lowered = name.lower()
    return any(marker in lowered for marker in _CONFLICTED_COPY_MARKERS)


def _should_skip(path: Path) -> str | None:
    """Return a skip reason string if this file should be skipped, else None."""
    name = path.name
    if name.startswith("."):
        return "dotfile"
    if name.startswith("~$"):
        return "office lock file"
    if _is_conflicted_copy(name):
        return "Dropbox conflicted copy"
    return None


def _iter_dir_entries(d: Path, recursive: bool):
    """Yield (Path, is_symlink) for files directly under `d` (and, if
    `recursive`, under its subdirectories), using `os.scandir` /
    `DirEntry.is_symlink()` / `is_file(follow_symlinks=False)` throughout so
    a symlinked file or directory is reported to the caller, never silently
    traversed into or read as if it were the real file.
    """
    try:
        entries = list(os.scandir(d))
    except OSError:
        return
    for entry in sorted(entries, key=lambda e: e.name):
        try:
            is_symlink = entry.is_symlink()
        except OSError:
            continue
        if is_symlink:
            yield Path(entry.path), True
            continue
        try:
            is_file = entry.is_file(follow_symlinks=False)
            is_dir = entry.is_dir(follow_symlinks=False)
        except OSError:
            continue
        if is_file:
            yield Path(entry.path), False
        elif is_dir and recursive:
            yield from _iter_dir_entries(Path(entry.path), recursive)


def extract_dir(
    d: Path,
    *,
    suffixes: set[str] | None = None,
    max_files: int | None = None,
    recursive: bool = False,
    max_bytes: int = 40 * 1024 * 1024,
    max_output_bytes: int = _DEFAULT_MAX_OUTPUT_BYTES,
    max_zip_uncompressed_bytes: int = _DEFAULT_MAX_ZIP_UNCOMPRESSED_BYTES,
) -> list[Extraction]:
    """Extract text from every eligible file directly under `d`.

    Non-recursive by default. Skips dotfiles, Office lock files (~$...), and
    Dropbox conflicted copies silently (these are not real content). Files
    over `max_bytes` are NOT silently dropped: an ok=False Extraction noting
    the size skip is appended so callers can see what was skipped.

    Symlinked entries (files or directories) are never followed: each is
    reported as an ok=False Extraction noting the refusal, same treatment as
    an oversized file, so a symlink planted inside a scanned directory
    cannot be used to pull content from outside it.
    """
    d = Path(d)
    results: list[Extraction] = []

    for path, is_symlink in _iter_dir_entries(d, recursive):
        if max_files is not None and len(results) >= max_files:
            break

        if is_symlink:
            results.append(_make_failed(path, "none", f"skipped: symlinked entry refused: {path}"))
            continue

        skip_reason = _should_skip(path)
        if skip_reason is not None:
            continue

        if suffixes is not None and path.suffix.lower() not in suffixes:
            continue

        try:
            size = path.stat().st_size
        except OSError as exc:
            results.append(_make_failed(path, "none", f"could not stat file: {exc}"))
            continue

        if size > max_bytes:
            results.append(_make_failed(path, "none", f"skipped: file size {size} bytes exceeds max_bytes={max_bytes}"))
            continue

        results.append(extract(
            path,
            max_output_bytes=max_output_bytes,
            max_zip_uncompressed_bytes=max_zip_uncompressed_bytes,
        ))

    return results


def cache_key(path: Path) -> str:
    """Deterministic cache key derived from realpath + size + mtime_ns.

    Note: this is stat-based, not content-hashed, so it is fast but relies on
    mtime changing when content changes (true for normal filesystem writes;
    an adversarial or clock-skewed write could in principle produce a
    collision — acceptable for this cache's purpose).
    """
    path = Path(path)
    real = os.path.realpath(path)
    st = os.stat(path)
    key_material = f"{real}|{st.st_size}|{st.st_mtime_ns}".encode("utf-8")
    return hashlib.sha256(key_material).hexdigest()


def _archive_root() -> Path:
    """Resolve the proposals archive root fresh from the environment (not
    cached at import time), so tests can swap AUTO_PROPOSALS_ROOT per case.
    Mirrors `lib.paths._default_root()`'s default; kept independent so this
    module has no import-time dependency on lib.paths.
    """
    return Path(os.environ.get("AUTO_PROPOSALS_ROOT", "~/Dropbox/Work/Proposals")).expanduser()


def extract_cached(path: Path, cache_dir: Path, *, timeout: int = 120,
                    max_output_bytes: int = _DEFAULT_MAX_OUTPUT_BYTES,
                    max_zip_uncompressed_bytes: int = _DEFAULT_MAX_ZIP_UNCOMPRESSED_BYTES) -> Extraction:
    """Like `extract`, but reads/writes a cache under `cache_dir`.

    `cache_dir` is required and must be a directory outside the proposals
    archive (e.g. ~/.cache/auto-proposals): this function never writes
    anywhere except directly under `cache_dir` (a `<key>.txt` + `<key>.json`
    sidecar pair), and never writes into `path`'s own directory.

    Both of those are hard refusals (`ExtractRefused`), not `assert`s: an
    `assert` disappears under `python -O`, and this module must be
    structurally incapable of writing into the archive regardless of how
    Python is invoked. The archive-containment check compares realpaths, so
    it also catches any nested subdirectory of the archive root, not just
    the root itself.
    """
    path = Path(path)
    cache_dir = Path(cache_dir)
    if not cache_dir:
        raise ExtractRefused("cache_dir is required")

    real_cache_dir = os.path.realpath(cache_dir)
    real_source_parent = os.path.realpath(path.parent)
    if real_cache_dir == real_source_parent:
        raise ExtractRefused("cache_dir must not be the source file's own directory")

    real_archive_root = os.path.realpath(_archive_root())
    if real_cache_dir == real_archive_root or real_cache_dir.startswith(real_archive_root + os.sep):
        raise ExtractRefused(
            f"cache_dir {cache_dir} resolves to {real_cache_dir!r}, which is inside the "
            f"proposals archive root {real_archive_root!r}; refusing to write there"
        )

    key = cache_key(path)
    text_path = cache_dir / f"{key}.txt"
    json_path = cache_dir / f"{key}.json"

    if text_path.exists() and json_path.exists():
        try:
            meta = json.loads(json_path.read_text(encoding="utf-8"))
            text = text_path.read_text(encoding="utf-8", errors="replace")
            return Extraction(
                path=Path(meta["path"]),
                text=text,
                tool=meta["tool"],
                ok=meta["ok"],
                note=meta["note"],
                chars=meta["chars"],
            )
        except (OSError, ValueError, KeyError):
            # Corrupt cache entry: fall through and regenerate.
            pass

    result = extract(
        path, timeout=timeout, max_output_bytes=max_output_bytes,
        max_zip_uncompressed_bytes=max_zip_uncompressed_bytes,
    )

    cache_dir.mkdir(parents=True, exist_ok=True)
    try:
        text_path.write_text(result.text, encoding="utf-8")
        json_path.write_text(
            json.dumps({
                "path": str(result.path),
                "tool": result.tool,
                "ok": result.ok,
                "note": result.note,
                "chars": result.chars,
            }),
            encoding="utf-8",
        )
    except OSError:
        # Cache write failed (e.g. disk full); still return the freshly
        # computed result rather than raising.
        pass

    return result


def summarise(extractions: list[Extraction]) -> str:
    """Compact plain-text table: file · tool · chars · ok · note, one row per line."""
    if not extractions:
        return "(no files)"

    rows = []
    for e in extractions:
        rows.append((str(e.path), e.tool, str(e.chars), "ok" if e.ok else "FAIL", e.note))

    widths = [max(len(row[i]) for row in rows) for i in range(4)]
    lines = []
    for row in rows:
        file_col, tool_col, chars_col, ok_col, note = row
        lines.append(
            f"{file_col.ljust(widths[0])}  {tool_col.ljust(widths[1])}  "
            f"{chars_col.rjust(widths[2])}  {ok_col.ljust(widths[3])}  {note}"
        )
    return "\n".join(lines)


def _main() -> None:
    import argparse

    parser = argparse.ArgumentParser(description="Extract text from proposal-archive documents.")
    parser.add_argument("path", type=Path, help="File or directory to extract from")
    parser.add_argument("--dir", action="store_true", help="Treat `path` as a directory")
    parser.add_argument("--recursive", action="store_true", help="Recurse into subdirectories (only with --dir)")
    parser.add_argument("--json", action="store_true", help="Emit JSON instead of plain text/table")
    args = parser.parse_args()

    if args.dir:
        extractions = extract_dir(args.path, recursive=args.recursive)
        if args.json:
            print(json.dumps([asdict(e) | {"path": str(e.path)} for e in extractions], ensure_ascii=False, indent=2))
        else:
            print(summarise(extractions))
    else:
        result = extract(args.path)
        if args.json:
            print(json.dumps(asdict(result) | {"path": str(result.path)}, ensure_ascii=False, indent=2))
        else:
            print(result.text)


if __name__ == "__main__":
    _main()
